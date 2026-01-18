
import os
import pandas as pd
from docx import Document
import openpyxl

SOURCE_DIR = 'original docs'
OUTPUT_BASE = os.path.join(SOURCE_DIR, 'markdown_docs')

def ensure_dir(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)

def extract_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def extract_xlsx(file_path):
    try:
        # Load workbook to get sheet names
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames
        
        full_text = []
        
        for sheet in sheet_names:
            full_text.append(f"## Sheet: {sheet}")
            df = pd.read_excel(file_path, sheet_name=sheet)
            # Convert to markdown table
            markdown_table = df.to_markdown(index=False)
            if markdown_table:
                full_text.append(markdown_table)
            else:
                full_text.append("(Empty Sheet)")
        
        return "\n\n".join(full_text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def process_file(file_path, relative_path):
    ext = os.path.splitext(file_path)[1].lower()
    content = ""
    
    if ext == '.docx':
        content = extract_docx(file_path)
    elif ext == '.xlsx':
        content = extract_xlsx(file_path)
    else:
        return
    
    # Create output path
    output_dir = os.path.join(OUTPUT_BASE, os.path.dirname(relative_path))
    ensure_dir(output_dir)
    
    output_filename = os.path.splitext(os.path.basename(file_path))[0] + ".md"
    output_path = os.path.join(output_dir, output_filename)
    
    # Create consistent markdown structure
    markdown_content = f"""# {os.path.basename(file_path)}

## Source Information
- **Original File**: `{relative_path}`
- **Type**: {ext}

## Content
{content}
"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"Docs Processed: {output_path}")

def main():
    ensure_dir(OUTPUT_BASE)
    
    for root, dirs, files in os.walk(SOURCE_DIR):
        # Skip output directory and hidden files
        if 'markdown_docs' in root:
            continue
            
        for file in files:
            if file.startswith('~$') or file.startswith('.'):
                continue
                
            if file.lower().endswith(('.docx', '.xlsx')):
                file_path = os.path.join(root, file)
                relative_path = os.path.relpath(file_path, SOURCE_DIR)
                print(f"Processing: {relative_path}")
                process_file(file_path, relative_path)

if __name__ == "__main__":
    main()
